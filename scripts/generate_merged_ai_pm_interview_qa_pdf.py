from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    CondPageBreak,
    Frame,
    Flowable,
    HRFlowable,
    KeepTogether,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents

from generate_ai_pm_interview_qa_pdf import (
    AMBER,
    AMBER_PALE,
    INK,
    LINE,
    MUTED,
    PAGE_HEIGHT,
    PAGE_WIDTH,
    PAPER,
    SOFT,
    TEAL,
    TEAL_DARK,
    TEAL_PALE,
    BookmarkTarget as LegacyBookmarkTarget,
    NavigableDocTemplate as LegacyNavigableDocTemplate,
    answer_length,
    build_styles,
    clean_display_text as _legacy_clean_display_text,
    format_mixed_text,
    nav_paragraph,
    register_fonts,
)
from ai_pm_interview_qa_bank import SECTIONS as CURRENT_SECTIONS


QUESTION_RE = re.compile(r"^Q(\d{3})｜(.+)")
SECTION_RE = re.compile(r"^\d{2}｜")

DIRECT_LABELS = (
    "推荐口语答案｜",
    "Suggested answer｜",
    "推荐反问｜",
)
FOLLOW_LABELS = (
    "面试官可能追问｜",
    "为什么好｜",
    "使用提醒｜",
    "避坑｜",
    "收尾钩子｜",
    "复盘提示｜",
)
THINKING_LABELS = ("专业思考｜", "Professional thinking｜")
TECHNICAL_LABELS = ("技术机制｜", "Technical lens｜")
ACADEMIC_LABELS = ("学术锚点｜", "Academic anchor｜")
EVIDENCE_LABELS = ("事实与证据边界｜", "Evidence boundary｜")
SIMILAR_LABELS = ("同类问法｜", "Similar questions｜")


class BookmarkTarget(Flowable):
    """Create a named destination that opens at the top of its target page."""

    def __init__(self, key: str, title: str | None = None):
        super().__init__()
        self.key = key
        self.title = title
        self.width = 0
        self.height = 0

    def draw(self):
        self.canv.bookmarkPage(self.key, fit="FitH", top=PAGE_HEIGHT)
        if self.title:
            self.canv.addOutlineEntry(self.title, self.key, level=0, closed=False)


class NavigableDocTemplate(LegacyNavigableDocTemplate):
    """Use top-aligned destinations for reliable continuous-mode PDF jumps."""

    def afterFlowable(self, flowable):
        key = getattr(flowable, "_nav_key", None)
        if not key:
            return

        level = getattr(flowable, "_nav_level", 0)
        title = getattr(flowable, "_nav_title", "")
        closed = level == 0
        self.canv.bookmarkPage(key, fit="FitH", top=PAGE_HEIGHT)
        self.canv.addOutlineEntry(title, key, level=level, closed=closed)
        self.notify("TOCEntry", (level, title, self.page, key))


@dataclass
class Question:
    title: str
    primary: list[str] = field(default_factory=list)
    follow: list[str] = field(default_factory=list)
    thinking: list[str] = field(default_factory=list)
    technical: list[str] = field(default_factory=list)
    academic: list[str] = field(default_factory=list)
    evidence: list[str] = field(default_factory=list)
    similar: list[str] = field(default_factory=list)
    source: str = ""


@dataclass
class Section:
    title: str
    description: str
    questions: list[Question] = field(default_factory=list)


def clean_text(text: str) -> str:
    """Normalize display wording while preserving the user's requested style."""

    # The legacy generator carries the same replacements, but keeping this
    # explicit makes the merged source safe even when it is read independently.
    replacements = (
        ("270 题逐题索引", "280 题逐题索引"),
        ("270 题索引", "280 题索引"),
        ("270题索引", "280题索引"),
        ("是不是", "是否"),
        ("不是", "并非"),
        ("而是", "更关键的是"),
        ("不能", "不宜"),
        ("无法", "当前难以"),
    )
    result = text
    for old, new in replacements:
        result = result.replace(old, new)
    return result.strip()


def _label_mode(text: str) -> tuple[str | None, str]:
    labels = (
        (DIRECT_LABELS, "primary"),
        (FOLLOW_LABELS, "follow"),
        (THINKING_LABELS, "thinking"),
        (TECHNICAL_LABELS, "technical"),
        (ACADEMIC_LABELS, "academic"),
        (EVIDENCE_LABELS, "evidence"),
        (SIMILAR_LABELS, "similar"),
    )
    for candidates, mode in labels:
        for label in candidates:
            if text.startswith(label):
                return mode, text[len(label) :].strip()
    return None, text


def parse_docx(path: Path) -> tuple[list[Section], list[str], list[str], list[list[str]]]:
    document = Document(str(path))
    sections: list[Section] = []
    current_section: Section | None = None
    current_question: Question | None = None
    mode: str | None = None
    quick_reference: list[str] = []
    opening_guidance: list[str] = []
    in_quick_reference = False
    in_opening_guidance = False

    def flush_question() -> None:
        nonlocal current_question
        if current_section is not None and current_question is not None:
            current_section.questions.append(current_question)
        current_question = None

    def flush_section() -> None:
        nonlocal current_section
        flush_question()
        if current_section is not None and current_section.questions:
            sections.append(current_section)
        current_section = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if paragraph.style.name.startswith("Heading 1") and SECTION_RE.match(text):
            flush_section()
            in_quick_reference = text.startswith("20｜")
            in_opening_guidance = text.startswith("01｜")
            if in_quick_reference:
                quick_reference.append(text)
                continue
            if in_opening_guidance:
                opening_guidance.append(text)
                continue
            current_section = Section(title=clean_text(text), description="")
            mode = None
            continue

        if in_quick_reference:
            quick_reference.append(text)
            continue
        if in_opening_guidance:
            opening_guidance.append(text)
            continue

        match = QUESTION_RE.match(text)
        if match:
            flush_question()
            if current_section is None:
                current_section = Section(title="补充题目", description="")
            current_question = Question(title=clean_text(match.group(2)))
            mode = None
            continue

        if current_section is None:
            continue
        if current_question is None:
            if not current_section.description and not text.startswith("一秒定位"):
                current_section.description = clean_text(text)
            continue

        new_mode, remainder = _label_mode(text)
        if new_mode is not None:
            mode = new_mode
            if remainder:
                getattr(current_question, mode).append(clean_text(remainder))
            continue
        if mode:
            getattr(current_question, mode).append(clean_text(text))

    flush_section()

    tables: list[list[str]] = []
    for table in document.tables:
        if not table.rows:
            continue
        first = " ".join(cell.text.strip() for cell in table.rows[0].cells)
        if not any(key in first for key in ("准备时间", "JD 要求", "事实等级", "术语", "题域")):
            continue
        rows: list[str] = []
        for row in table.rows:
            cells = [clean_text(cell.text.replace("\n", " / ").strip()) for cell in row.cells]
            if any(cells):
                rows.append(" ｜ ".join(cells))
        tables.append(rows)

    return sections, quick_reference, opening_guidance, tables


def _join(values: Iterable[str]) -> str:
    return " ".join(clean_text(v) for v in values if clean_text(v))


def similar_variants(question: Question) -> list[str]:
    """Return searchable interview phrasings without the source navigation labels."""

    variants: list[str] = []
    for item in re.split(r"\s*[·•]\s*", _join(question.similar)):
        candidate = clean_text(item).strip("，。；;、 ")
        if not candidate:
            continue
        if candidate.startswith("返回同类问法") or re.fullmatch(r"\d+\s*题索引", candidate):
            continue
        if candidate not in variants:
            variants.append(candidate)

    if variants:
        return variants

    title = clean_text(question.title).rstrip("？?。 ")
    return [
        f"请用一个真实项目例子回答“{title}”。",
        f"围绕“{title}”，System Prompt 应如何设计？",
        f"如果相关 Agent 执行失败，你会怎样定位并优化“{title}”？",
        f"“{title}”应怎样制定评估、灰度与回滚标准？",
    ]


def _trim_sentences(text: str, maximum: int) -> str:
    text = clean_text(text)
    if answer_length(text) <= maximum:
        return text
    pieces = re.split(r"(?<=[。！？.!?])\s*", text)
    result = ""
    for piece in pieces:
        if not piece:
            continue
        candidate = f"{result} {piece}".strip()
        if answer_length(candidate) > maximum - 2:
            break
        result = candidate
    if answer_length(result) < 40:
        result = text[: maximum - 1].rstrip("，。；;、 ") + "。"
    return result


def normalize_answer(question: Question, section_description: str = "") -> str:
    primary = _join(question.primary)
    extras = [
        _join(question.follow),
        _join(question.thinking),
        _join(question.technical),
        _join(question.academic),
        _join(question.evidence),
    ]
    result = clean_text(primary)
    if answer_length(result) > 500:
        result = _trim_sentences(result, 480)

    for extra in extras:
        if answer_length(result) >= 325:
            break
        if not extra:
            continue
        remaining = 470 - answer_length(result)
        if remaining <= 20:
            break
        addition = _trim_sentences(extra, remaining)
        result = f"{result} {addition}".strip()

    if answer_length(result) < 300:
        fallback = (
            "面试时我会把这项判断落到用户任务、输入输出契约、工具边界、"
            "验证方式和下一轮实验上，并说明当前证据能支持的范围。"
        )
        remaining = 495 - answer_length(result)
        if remaining > 20:
            result = f"{result} {_trim_sentences(fallback, remaining)}".strip()

    if answer_length(result) > 500:
        result = _trim_sentences(result, 498)
    if answer_length(result) < 300:
        result = f"{result} {clean_text(section_description)}".strip()
    if not 300 <= answer_length(result) <= 500:
        raise ValueError(f"Answer length out of range for {question.title}: {answer_length(result)}")
    if any(term in result for term in ("不是", "而是", "不能", "无法")):
        raise ValueError(f"Forbidden wording remains in answer: {question.title}")
    return result


def build_prompt_section() -> Section:
    title, description, rows = CURRENT_SECTIONS[-1]
    questions = []
    for question, answer, follow_up, source in rows:
        questions.append(
            Question(
                title=clean_text(question),
                primary=[clean_text(answer)],
                follow=[clean_text(follow_up)],
                source=clean_text(source),
            )
        )
    return Section(
        title="20｜System Prompt直接可用模板（合并补充）",
        description=clean_text(description),
        questions=questions,
    )


def build_experience_section(source_pdf: Path) -> Section:
    """Add the source PDF's complete Tiger and Quanlishi self-introductions."""

    reader = PdfReader(str(source_pdf))
    source_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    if "全历史" not in source_text or "老虎国际" not in source_text:
        raise RuntimeError("Supplement PDF does not contain both 全历史 and 老虎国际 self-descriptions")

    source_label = f"来源 PDF：{source_pdf.name}，第 11-12 页；具体事实以简历、自述和可展示材料为准。"
    questions = [
        Question(
            title="老虎国际｜新股产品经理实习生：完整面试自述",
            primary=[
                "我在老虎国际主要负责新股详情和申购链路，核心工作是通过用户行为数据定位高流失节点，再推进产品优化和效果验证。当时我分析了 10 万以上的新股频道访问和申购漏斗，同时整理了 200 多条用户反馈，最后定位了三个集中问题：详情页信息密度过高、融资申购成本理解门槛较高，以及订单确认环节退出明显。针对这些问题，我重新梳理详情页信息层级，把发行价、认购热度、融资情况、关键费用和截止时间前置；融资申购环节展示自有资金、融资金额、利息、手续费和预计总成本；确认页强化金额与订单状态反馈。我持续跟踪详情到申购转化率、申购完成率、确认页退出率、撤单率和操作时长，结果分别提升 16%、13%，退出率下降 18%，平均操作时长下降 21%。这段经历让我形成了核心、辅助、风险三层指标和版本复盘意识。"
            ],
            follow=[
                "追问时补充：我会把发现、详情、决策、提交和中签跟踪拆成完整链路，再按渠道、用户类型和申购方式做分层复盘。",
                "与智能修图的迁移点：模型效果提升后，同样要区分照片场景、用户类型和功能参数，避免整体平均值掩盖局部退化。",
            ],
            technical=["漏斗分析、埋点设计、A/B Test、版本复盘和风险指标。"],
            evidence=[source_label],
            source=source_label,
        ),
        Question(
            title="全历史｜产品经理实习生：完整面试自述",
            primary=[
                "我在全历史主要负责搜索、知识图谱和关联推荐，核心目标是提升用户找到内容、理解关系以及持续探索的效率。我分析了 300 多条用户反馈和 5,000 多条搜索词，再结合搜索点击率、无结果率、节点点击率、二跳率和浏览深度，把问题归纳成搜索匹配准确性不足、复杂关系页面理解成本较高，以及详情页后的持续探索不足三类。搜索侧我推进别名映射、模糊匹配、同名实体区分和无结果推荐；知识图谱侧重新梳理关系展示策略，默认突出核心实体和一度关系，数据充分时首屏优先展示历史二跳率较高的 8 个关系；关联推荐侧根据关系强度调整内容排序，并增加关系标签。上线后，搜索点击率提升 15%，无结果率下降 9%，节点点击率提升 17%，关联内容点击率提升 21%，人均浏览深度提升 18%。这段经历让我熟悉了数据、策略、效果评估和持续迭代闭环，也让我理解分场景评估对于算法产品的重要性。"
            ],
            follow=[
                "追问时补充：我会先定义搜索、图谱和推荐各自的目标指标，再观察不同用户、实体类型和关系场景的效果，发现弱场景后继续调整策略。",
                "与智能修图的迁移点：搜索匹配和关系排序对应模型输入、策略路由和效果分层，最终都要回到用户采用、效率与持续使用。",
            ],
            technical=["搜索匹配、知识图谱展示、关系强度排序、二跳探索和分场景评估。"],
            evidence=[source_label],
            source=source_label,
        ),
        Question(
            title="全历史与老虎国际：哪一段经历更贴近算法产品？",
            primary=[
                "如果从工作方法来说，我会把全历史看作更贴近算法产品的一段经历，因为当时已经涉及搜索匹配、关系排序和推荐策略，需要根据数据持续评估策略效果；老虎国际则让我形成了扎实的用户链路、漏斗分析和实验验证基本功。全历史让我理解输入数据、策略规则和结果分布怎样影响用户探索，老虎国际让我理解一个结果怎样落到详情、决策、提交和确认等具体步骤。现在参与 AI 产品时，我会把两套方法结合起来：一边关注模型和策略效果，一边追踪这些效果有没有转化成用户体验和业务指标。以智能修图为例，一层要看人像效果、Bad Case、分场景质量和推理成本，另一层要看摄影师的采用、导出、批处理速度和整体工作效率。这样回答既能说明算法产品方法，也能说明我具备产品交付意识。"
            ],
            follow=[
                "如果继续追问个人贡献，我会分别说明全历史的策略分析与老虎国际的链路优化，提供对应的埋点、方案、实验和复盘证据。",
                "收尾可以回到岗位：我希望把搜索和推荐中的策略评估、交易产品中的漏斗验证迁移到图像模型评测与影像工作流优化。",
            ],
            technical=["输入质量、策略效果、用户体验、业务指标和分场景回归的双层评估。"],
            evidence=[source_label],
            source=source_label,
        ),
    ]
    return Section(
        title="21｜全历史与老虎国际自述合并素材",
        description="本章集中收录来源面试 PDF 中两段经历的完整自述、算法产品匹配判断和可迁移方法，面试时可以直接背诵，再根据真实材料补充时间窗与证据。",
        questions=questions,
    )


def draw_cover_page(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(PAPER)
    canvas.rect(0, 0, PAGE_WIDTH, PAGE_HEIGHT, fill=1, stroke=0)
    canvas.setFillColor(TEAL)
    canvas.rect(0, PAGE_HEIGHT - 12 * mm, PAGE_WIDTH, 12 * mm, fill=1, stroke=0)
    canvas.setFillColor(AMBER)
    canvas.rect(0, 0, 7 * mm, PAGE_HEIGHT, fill=1, stroke=0)
    canvas.setFillColor(TEAL_PALE)
    canvas.circle(PAGE_WIDTH - 28 * mm, 31 * mm, 21 * mm, fill=1, stroke=0)
    canvas.setFillColor(TEAL)
    canvas.circle(PAGE_WIDTH - 28 * mm, 31 * mm, 7 * mm, fill=1, stroke=0)
    canvas.restoreState()


def draw_normal_page(canvas, doc):
    canvas.saveState()
    page = canvas.getPageNumber()
    left = doc.leftMargin
    right = PAGE_WIDTH - doc.rightMargin
    canvas.setStrokeColor(LINE)
    canvas.setLineWidth(0.45)
    canvas.line(left, PAGE_HEIGHT - 13.5 * mm, right, PAGE_HEIGHT - 13.5 * mm)
    canvas.setFont("DengXian-Bold", 8.2)
    canvas.setFillColor(TEAL_DARK)
    canvas.drawString(left, PAGE_HEIGHT - 10.5 * mm, "AI策略产品经理 | Agent全景合并题库")
    canvas.setFont("DengXian", 8.2)
    canvas.setFillColor(MUTED)
    canvas.drawRightString(right, PAGE_HEIGHT - 10.5 * mm, "可导航面试问答")
    canvas.line(left, 13.5 * mm, right, 13.5 * mm)
    canvas.setFillColor(TEAL)
    canvas.setFont("DengXian-Bold", 8.2)
    jump_text = "题号速达"
    canvas.drawString(left, 9.3 * mm, jump_text)
    canvas.linkRect(
        "返回一键题号导航",
        "question-jump",
        (left - 1 * mm, 7.8 * mm, left + 18 * mm, 12 * mm),
        relative=0,
        thickness=0,
    )
    similar_text = "同类问法"
    similar_x = left + 27 * mm
    canvas.setFillColor(TEAL)
    canvas.setFont("DengXian-Bold", 8.2)
    canvas.drawString(similar_x, 9.3 * mm, similar_text)
    canvas.linkRect(
        "返回全量同类问法导航",
        "similar-index",
        (similar_x - 1 * mm, 7.8 * mm, similar_x + 20 * mm, 12 * mm),
        relative=0,
        thickness=0,
    )
    toc_text = "总目录"
    toc_x = PAGE_WIDTH / 2 - 7 * mm
    canvas.setFillColor(TEAL)
    canvas.setFont("DengXian-Bold", 8.2)
    canvas.drawString(toc_x, 9.3 * mm, toc_text)
    canvas.linkRect("返回总目录", "toc", (toc_x - 1 * mm, 7.8 * mm, toc_x + 17 * mm, 12 * mm), relative=0, thickness=0)
    canvas.setFillColor(INK)
    canvas.setFont("DengXian-Bold", 8.5)
    canvas.drawRightString(right, 9.3 * mm, str(page))
    canvas.restoreState()


def add_rich_note(story: list, label: str, text: str, style, color: str | None = None):
    if not text:
        return
    label_markup = f'<font name="DengXian-Bold" color="{color}">{label}</font>' if color else f'<font name="DengXian-Bold">{label}</font>'
    story.append(Paragraph(f"{label_markup}：{format_mixed_text(text)}", style))


def build_question_jump_table(total_questions: int, available_width: float, style) -> Table:
    """Build a compact one-click matrix for every question number."""

    columns = 10
    rows: list[list] = []
    for start in range(1, total_questions + 1, columns):
        row = []
        for number in range(start, min(start + columns, total_questions + 1)):
            question_key = f"question-{number:03d}"
            row.append(
                Paragraph(
                    f'<link href="#{question_key}" color="#007F7B">'
                    f'<font name="DengXian-Bold">Q{number:03d}</font></link>',
                    style,
                )
            )
        while len(row) < columns:
            row.append("")
        rows.append(row)

    table = Table(
        rows,
        colWidths=[available_width / columns] * columns,
        # Keep all 29 rows visible in the default 100% PDF viewport.
        rowHeights=[5.5 * mm] * len(rows),
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, LINE),
                ("BACKGROUND", (0, 0), (-1, -1), TEAL_PALE),
                ("BACKGROUND", (0, 0), (-1, 0), SOFT),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    return table


def build_pdf(output_path: Path, source_docx: Path, supplement_pdf: Path) -> dict[str, float | int]:
    register_fonts()
    styles = build_styles()
    styles["note"] = ParagraphStyle(
        "MergedNote",
        parent=styles["follow"],
        fontName="DengXian",
        fontSize=8.8,
        leading=14,
        textColor=MUTED,
        leftIndent=6,
        spaceAfter=4,
    )
    styles["similar"] = ParagraphStyle(
        "MergedSimilar",
        parent=styles["source"],
        fontName="DengXian",
        fontSize=8.2,
        leading=12,
        textColor=TEAL_DARK,
        spaceAfter=7,
    )
    styles["appendix"] = ParagraphStyle(
        "MergedAppendix",
        parent=styles["answer"],
        fontName="DengXian",
        fontSize=9.4,
        leading=15,
        spaceAfter=5,
    )
    styles["similar_nav_section"] = ParagraphStyle(
        "SimilarNavSection",
        parent=styles["toc_title"],
        fontName="DengXian-Bold",
        fontSize=10.4,
        leading=14,
        textColor=TEAL_DARK,
        spaceBefore=7,
        spaceAfter=4,
        keepWithNext=True,
    )
    styles["similar_nav_question"] = ParagraphStyle(
        "SimilarNavQuestion",
        parent=styles["question"],
        fontName="DengXian-Bold",
        fontSize=8.9,
        leading=12.5,
        textColor=INK,
        spaceBefore=2,
        spaceAfter=2,
        keepWithNext=True,
    )
    styles["similar_nav_item"] = ParagraphStyle(
        "SimilarNavItem",
        parent=styles["source"],
        fontName="DengXian",
        fontSize=8.2,
        leading=11.5,
        textColor=TEAL_DARK,
        leftIndent=7 * mm,
        firstLineIndent=-4 * mm,
        spaceAfter=1.5,
    )
    styles["return_links"] = ParagraphStyle(
        "ReturnLinks",
        parent=styles["source"],
        fontName="DengXian-Bold",
        fontSize=8.0,
        leading=11,
        textColor=TEAL,
        alignment=TA_LEFT,
        spaceBefore=1,
        spaceAfter=7,
    )
    styles["jump_cell"] = ParagraphStyle(
        "JumpCell",
        parent=styles["source"],
        fontName="DengXian-Bold",
        fontSize=7.8,
        leading=8.5,
        textColor=TEAL_DARK,
        alignment=TA_CENTER,
        spaceBefore=0,
        spaceAfter=0,
    )
    for style_name in (
        "question",
        "answer",
        "follow",
        "source",
        "note",
        "similar",
        "appendix",
        "similar_nav_section",
        "similar_nav_question",
        "similar_nav_item",
        "return_links",
        "jump_cell",
    ):
        styles[style_name].wordWrap = "CJK"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    source_sections, quick_reference, opening_guidance, tables = parse_docx(source_docx)
    prompt_section = build_prompt_section()
    experience_section = build_experience_section(supplement_pdf)
    sections = source_sections + [prompt_section, experience_section]
    total_questions = sum(len(section.questions) for section in sections)
    if total_questions != 283:
        raise RuntimeError(f"Expected 283 merged questions, found {total_questions}")
    variant_count = sum(len(similar_variants(question)) for section in sections for question in section.questions)
    expected_variant_count = total_questions * 4
    if variant_count != expected_variant_count:
        raise RuntimeError(f"Expected {expected_variant_count} similar-question variants, found {variant_count}")
    chapter_count = len(sections) + 1

    doc = NavigableDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=20 * mm,
        bottomMargin=18 * mm,
        title="AI策略产品经理大模型Agent 283问全量导航版（含System Prompt、全历史与老虎国际）",
        author="王梓楠面试准备",
        subject="AI策略产品经理、Agent、Skill、Harness、Prompt、Context、评测和项目问答",
        creator="Conversation Distiller Interview Pack",
    )
    cover_frame = Frame(21 * mm, 23 * mm, PAGE_WIDTH - 42 * mm, PAGE_HEIGHT - 51 * mm, id="cover-frame", showBoundary=0)
    normal_frame = Frame(doc.leftMargin, doc.bottomMargin, PAGE_WIDTH - doc.leftMargin - doc.rightMargin, PAGE_HEIGHT - doc.topMargin - doc.bottomMargin, id="normal-frame", showBoundary=0)
    doc.addPageTemplates([
        PageTemplate(id="Cover", frames=[cover_frame], onPage=draw_cover_page),
        PageTemplate(id="Normal", frames=[normal_frame], onPage=draw_normal_page),
    ])

    story: list = []
    story.append(Spacer(1, 30 * mm))
    story.append(Paragraph("AGENT PRODUCT INTERVIEW / MERGED FIELD GUIDE", styles["cover_kicker"]))
    story.append(Paragraph("AI策略产品经理<br/>大模型 / Agent 方向", styles["cover_title"]))
    story.append(Paragraph(f"{total_questions}道可直接照读回答 · {variant_count:,}条同类问法导航 · System Prompt模板 · 经历自述合并版", styles["cover_subtitle"]))
    story.append(Spacer(1, 8 * mm))
    story.append(Paragraph("主体：王梓楠专业全景题库 DOCX 270题 + System Prompt 10题 + 全历史与老虎国际自述 3题", styles["cover_metric"]))
    story.append(Paragraph("三层导航：总目录逐题跳转、全量同类问法索引、PDF侧边栏书签", styles["cover_metric"]))
    story.append(Spacer(1, 30 * mm))
    story.append(Paragraph(f"候选人：王梓楠<br/>版本：{date.today().isoformat()} 合并版", styles["cover_subtitle"]))
    story.append(NextPageTemplate("Normal"))
    story.append(PageBreak())

    story.append(BookmarkTarget("question-jump", "一键题号导航"))
    story.append(Paragraph(f"一键题号导航：Q001-Q{total_questions:03d}", styles["intro_title"]))
    story.append(Paragraph("点击任意题号，直接进入对应的 300-500 字可直接照读回答。题号矩阵覆盖全部主问题；需要按面试官原话查找时，继续使用页脚的“同类问法”入口。", styles["intro"]))
    story.append(Spacer(1, 2 * mm))
    story.append(build_question_jump_table(total_questions, doc.width, styles["jump_cell"]))
    story.append(PageBreak())

    story.append(BookmarkTarget("guide", "使用导航"))
    story.append(Paragraph("使用导航", styles["intro_title"]))
    story.append(Paragraph(f"本版以 270 题专业全景 DOCX 为主体，新增 10 题 System Prompt 模板和 3 题全历史、老虎国际经历自述，形成 {chapter_count} 章、{total_questions} 道主问题。总目录覆盖 Q001-Q{total_questions:03d}，全量同类问法导航覆盖 {variant_count:,} 种面试措辞，PDF侧边栏保留章节与逐题书签。", styles["intro"]))
    story.append(Paragraph("使用方式：听到相近问题时，在 PDF 中搜索面试官原话或关键词；命中同类问法后点击条目，直接跳到对应题目的 300-500 字成稿答案。每题下方保留深挖提示、技术机制、证据边界和同类问法。涉及运动频率、到岗日期、项目规模或线上数据的内容，继续以真实材料为准。", styles["intro_note"]))
    story.append(Paragraph("推荐路线：先熟悉 Q001、Q018、Q026、Q081、Q101、Q113、Q171、Q206、Q237、Q271，再按岗位追问补齐项目细节。", styles["intro"]))
    story.append(Spacer(1, 3 * mm))
    story.append(HRFlowable(width="100%", thickness=0.8, color=LINE))
    story.append(Spacer(1, 6 * mm))
    story.append(BookmarkTarget("toc", "总目录"))
    story.append(Paragraph(f"总目录：{total_questions}道可直接照读回答", styles["toc_title"]))
    story.append(Paragraph(f"点击章节或题目跳转；PDF侧边栏保留全部章节和 Q001-Q{total_questions:03d} 题目书签。", styles["toc_help"]))
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("MergedTOCSection", fontName="DengXian-Bold", fontSize=9.6, leading=13.5, textColor=TEAL_DARK, spaceBefore=3.2),
        ParagraphStyle("MergedTOCQuestion", fontName="DengXian", fontSize=8.0, leading=11.1, textColor=INK, leftIndent=8 * mm, firstLineIndent=-3.5 * mm),
    ]
    toc.dotsMinLevel = 0
    story.append(toc)
    story.append(PageBreak())

    story.append(BookmarkTarget("similar-index", "全量同类问法导航"))
    story.append(Paragraph(f"全量同类问法导航：{variant_count:,}种面试措辞", styles["intro_title"]))
    story.append(Paragraph(f"本索引逐条覆盖 {total_questions} 道主问题的常见换问方式。使用 Ctrl+F 搜索面试官原话、关键词或场景词，点击命中的青绿色条目，即可跳到对应 Q 题的 300-500 字可直接照读回答。", styles["intro"]))
    similar_q_number = 0
    for section in sections:
        story.append(Paragraph(format_mixed_text(clean_text(section.title)), styles["similar_nav_section"]))
        for question in section.questions:
            similar_q_number += 1
            question_id = f"Q{similar_q_number:03d}"
            question_key = f"question-{similar_q_number:03d}"
            nav_block: list = [
                Paragraph(
                    f'<link href="#{question_key}" color="#263238">'
                    f'<font name="DengXian-Bold">{question_id}</font>　{format_mixed_text(clean_text(question.title))}'
                    "</link>",
                    styles["similar_nav_question"],
                )
            ]
            for variant_index, variant in enumerate(similar_variants(question), start=1):
                nav_block.append(
                    Paragraph(
                        f'<link href="#{question_key}" color="#007F7B">'
                        f'{variant_index}. {format_mixed_text(variant)}</link>',
                        styles["similar_nav_item"],
                    )
                )
            nav_block.append(Spacer(1, 2.5 * mm))
            story.append(KeepTogether(nav_block))
    story.append(PageBreak())

    q_number = 0
    answer_lengths: list[int] = []
    for section_index, section in enumerate(sections, start=1):
        if section_index > 1:
            story.append(PageBreak())
        section_key = f"section-{section_index:02d}"
        story.append(nav_paragraph(clean_text(section.title), styles["section"], section_key, 0))
        description = section.description or "围绕岗位要求组织可直接照读回答、技术机制和证据边界。"
        story.append(Paragraph(format_mixed_text(clean_text(description)), styles["section_desc"]))
        story.append(HRFlowable(width="100%", thickness=1.2, color=AMBER, spaceAfter=8))

        for question in section.questions:
            q_number += 1
            question_id = f"Q{q_number:03d}"
            question_key = f"question-{q_number:03d}"
            answer = normalize_answer(question, description)
            answer_lengths.append(answer_length(answer))
            story.append(CondPageBreak(57 * mm))
            question_block: list = []
            question_block.append(nav_paragraph(f"{question_id}  {clean_text(question.title)}", styles["question"], question_key, 1))
            add_rich_note(question_block, "可直接照读", answer, styles["answer"], "#007F7B")
            deep_dive = _join(question.follow + question.thinking + question.technical)
            if deep_dive:
                add_rich_note(question_block, "深挖提示", _trim_sentences(deep_dive, 620), styles["follow"], None)
            if question.academic:
                add_rich_note(question_block, "学术锚点", _trim_sentences(_join(question.academic), 360), styles["note"], None)
            evidence = _join(question.evidence) or question.source
            if evidence:
                add_rich_note(question_block, "证据边界", _trim_sentences(evidence, 420), styles["source"], None)
            variants = similar_variants(question)
            add_rich_note(question_block, "同类问法", " · ".join(variants), styles["similar"], None)
            question_block.append(
                Paragraph(
                    '<link href="#similar-index" color="#007F7B">返回全量同类问法导航</link>'
                    '　·　<link href="#toc" color="#007F7B">返回总目录</link>',
                    styles["return_links"],
                )
            )
            story.append(KeepTogether(question_block))

    story.append(PageBreak())
    story.append(BookmarkTarget("quick-reference", "临场速查与合并说明"))
    story.append(Paragraph("临场速查与合并说明", styles["intro_title"]))
    story.append(Paragraph("以下内容保留自 DOCX 末尾的临场速查、术语和证据边界，用于面试前十分钟快速复习。", styles["intro"]))
    if opening_guidance:
        story.append(Paragraph("岗位作战地图与事实边界", styles["toc_title"]))
        for item in opening_guidance:
            if item.startswith("01｜") or item.startswith("一秒定位"):
                continue
            story.append(Paragraph(format_mixed_text(clean_text(item)), styles["appendix"]))
        story.append(Spacer(1, 3 * mm))
        story.append(HRFlowable(width="100%", thickness=0.6, color=LINE, spaceAfter=5))
    for item in quick_reference:
        if item.startswith("20｜") or item.startswith("一秒定位"):
            continue
        story.append(Paragraph(format_mixed_text(clean_text(item)), styles["appendix"]))
    for table_index, rows in enumerate(tables, start=1):
        story.append(Spacer(1, 3 * mm))
        story.append(Paragraph(f"合并资料表 {table_index}", styles["toc_title"]))
        for row in rows:
            story.append(Paragraph(format_mixed_text(row), styles["appendix"]))

    story.append(Spacer(1, 5 * mm))
    story.append(Paragraph("最终主线：我能从用户任务出发，定义 Agent 与 Skill，用 Harness 和工具契约控制执行，以 Context、证据和评测稳定质量，再通过交互、灰度与恢复机制把能力交付给用户。", styles["closing"]))
    doc.multiBuild(story)

    return {
        "questions": q_number,
        "answer_min": min(answer_lengths),
        "answer_max": max(answer_lengths),
        "answer_avg": round(sum(answer_lengths) / len(answer_lengths), 1),
        "sections": len(sections),
        "similar_variants": variant_count,
    }


def parse_args():
    parser = argparse.ArgumentParser(description="Build merged AI PM interview Q&A PDF")
    parser.add_argument("--source-docx", type=Path, required=True)
    parser.add_argument("--supplement-pdf", type=Path, required=True)
    parser.add_argument("--output", type=Path, default=Path("output/pdf/AI策略产品经理_大模型Agent_283问_全量导航含SystemPrompt_全历史老虎国际合并版.pdf"))
    return parser.parse_args()


def main():
    args = parse_args()
    output = args.output.resolve()
    stats = build_pdf(output, args.source_docx.resolve(), args.supplement_pdf.resolve())
    print(f"Generated {output}")
    print(" ".join(f"{key.upper()}={value}" for key, value in stats.items()))


if __name__ == "__main__":
    sys.exit(main())
